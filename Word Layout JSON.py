from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json


def extract_docx_layout_to_json(docx_path, out_path, page_limit=1):
    doc = Document(docx_path)
    layout = {"tables": []}

    for table_idx, table in enumerate(doc.tables):
        if table_idx >= page_limit:
            break

        table_data = {"columns": [], "rows": []}

        max_cols = max(len(row.cells) for row in table.rows)

        for col_idx in range(max_cols):
            table_data["columns"].append(
                {"name": f"Column {col_idx + 1}", "width": None, "x": col_idx}
            )

        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_data = {"text": cell.text.strip(), "style": {}}

                if cell.paragraphs:
                    p = cell.paragraphs[0]
                    run = p.runs[0] if p.runs else None
                    align = p.alignment

                    if align == WD_ALIGN_PARAGRAPH.LEFT:
                        align_str = "LEFT"
                    elif align == WD_ALIGN_PARAGRAPH.CENTER:
                        align_str = "CENTER"
                    elif align == WD_ALIGN_PARAGRAPH.RIGHT:
                        align_str = "RIGHT"
                    elif align == WD_ALIGN_PARAGRAPH.JUSTIFY:
                        align_str = "JUSTIFY"
                    else:
                        align_str = "LEFT"

                    cell_data["style"] = {
                        "bold": run.bold if run else False,
                        "italic": run.italic if run else False,
                        "underline": run.underline if run else False,
                        "font": run.font.name if run and run.font.name else None,
                        "size": run.font.size.pt if run and run.font.size else None,
                        "alignment": align_str,
                    }

                row_data.append(cell_data)
            table_data["rows"].append(row_data)

        layout["tables"].append(table_data)

    with open(out_path, "w") as f:
        json.dump(layout, f, indent=2)


# Usage
extract_docx_layout_to_json("Openground log3.docx", "borehole_layout_page1.json")
