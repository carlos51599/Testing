from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
import json


def get_cell_width(cell):
    tc = cell._tc
    tc_pr = tc.tcPr
    width_elem = tc_pr.find(qn("w:tcW"))
    if width_elem is not None:
        return int(width_elem.get(qn("w:w")))
    return None


def get_grid_span(cell):
    span_elem = cell._tc.tcPr.find(qn("w:gridSpan"))
    if span_elem is not None:
        return int(span_elem.get(qn("w:val")))
    return 1


def get_vmerge(cell):
    vmerge_elem = cell._tc.tcPr.find(qn("w:vMerge"))
    if vmerge_elem is not None:
        return vmerge_elem.get(qn("w:val")) or "continue"
    return None


def extract_cell_format(cell):
    content = []
    for para in cell.paragraphs:
        para_runs = []
        for run in para.runs:
            font = run.font
            para_runs.append(
                {
                    "text": run.text,
                    "bold": font.bold,
                    "italic": font.italic,
                    "underline": font.underline,
                    "font_name": font.name,
                    "font_size": font.size.pt if font.size else None,
                }
            )
        content.append(para_runs)
    return content


def extract_tables(doc_path):
    doc = Document(doc_path)
    tables_json = []

    for t_idx, table in enumerate(doc.tables):
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(
                    {
                        "content": extract_cell_format(cell),
                        "gridSpan": get_grid_span(cell),
                        "vMerge": get_vmerge(cell),
                        "width": get_cell_width(cell),
                    }
                )
            table_data.append(row_data)
        tables_json.append({"table_index": t_idx, "rows": table_data})

    with open("rich_layout_tables.json", "w") as f:
        json.dump(tables_json, f, indent=2)


extract_tables("Openground log3.docx")
